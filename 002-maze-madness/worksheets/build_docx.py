#!/usr/bin/env python
"""Render the Maze Madness worksheet markdown into clean, print-ready .docx.

Goals (matching the PDF look as closely as Word allows):
  * ruled "writing space" boxes wherever the source has <div class="write-space">
  * every heading starts on a new page
  * a question and its answer box never get split across two pages
"""
import glob
import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

LINE_PX = 28  # one ruled line in the source CSS

INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_runs(p, text):
    """Add inline runs handling **bold** and `code`."""
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10.5)
            r._element.rPr.rFonts.set(qn("w:cs"), "Consolas")
            shade(r, "EEEEEE")
        else:
            p.add_run(tok)


def shade(run, hexcolor):
    rpr = run._element.get_or_add_rPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    rpr.append(sh)


def set_keep(p, with_next):
    pf = p.paragraph_format
    pf.keep_together = True       # never split this paragraph's lines
    pf.keep_with_next = with_next  # glue to the following paragraph


def bottom_border(p, color="C9C9C9"):
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)


def code_shading(p, hexcolor="F4F4F4"):
    ppr = p._p.get_or_add_pPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    ppr.append(sh)


def render(md_path, out_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()

    # base styling
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    for sec in doc.sections:
        sec.top_margin = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)

    paras = []  # collect (paragraph, is_boxline) so we can fix keep_with_next after

    def emit(p, box_end=False):
        # box_end marks the last ruled line of an answer box: a break is allowed after it
        paras.append((p, box_end))

    first_heading_seen = False
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # fenced code block
        if line.startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # closing fence
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(8)
            code_shading(p)
            r = p.add_run("\n".join(buf))
            r.font.name = "Consolas"
            r.font.size = Pt(10.5)
            r._element.rPr.rFonts.set(qn("w:cs"), "Consolas")
            emit(p)
            continue

        # writing-space box
        m = re.search(r'class="write-space([^"]*)"(?:[^>]*min-height:\s*(\d+))?', line)
        if line.strip().startswith("<div") and "write-space" in line:
            cls, px = m.group(1), m.group(2)
            if px:
                height = int(px)
            elif "tall" in cls:
                height = 170
            elif "short" in cls:
                height = 86
            else:
                height = 114
            nlines = max(2, round(height / LINE_PX))
            for k in range(nlines):
                bp = doc.add_paragraph()
                bp.paragraph_format.space_after = Pt(0)
                bp.paragraph_format.space_before = Pt(0)
                bp.paragraph_format.line_spacing = Pt(26)
                bottom_border(bp)
                bp.add_run(" ")
                emit(bp, box_end=(k == nlines - 1))
            i += 1
            continue

        stripped = line.strip()

        # blank / horizontal rule -> skip (page breaks + spacing handle separation)
        if stripped == "" or re.fullmatch(r"-{3,}", stripped):
            i += 1
            continue

        # headings
        if stripped.startswith("#"):
            hm = re.match(r"(#{1,4})\s+(.*)", stripped)
            level = len(hm.group(1))
            text = hm.group(2)
            p = doc.add_paragraph()
            if level == 1:
                p.paragraph_format.space_after = Pt(4)
                add_runs(p, text)
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(20)
            else:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(6)
                add_runs(p, text)
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(15 if level == 2 else 12.5)
                    r.font.color.rgb = RGBColor(0x1A, 0x4D, 0x2E)
                # new page before every heading (except the very first H1)
                if first_heading_seen:
                    p.paragraph_format.page_break_before = True
            if level <= 2:
                first_heading_seen = True
            emit(p)
            i += 1
            continue

        # blockquote (may contain nested list items / blank > lines)
        if stripped.startswith(">"):
            while i < n and lines[i].strip().startswith(">"):
                content = re.sub(r"^\s*>\s?", "", lines[i])
                if content.strip() == "":
                    i += 1
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.35)
                p.paragraph_format.space_after = Pt(3)
                lm = re.match(r"(\d+)\.\s+(.*)", content) or re.match(r"-\s+(.*)", content)
                if content.lstrip().startswith("- "):
                    p.paragraph_format.left_indent = Inches(0.6)
                    add_runs(p, "•  " + content.lstrip()[2:])
                elif re.match(r"\d+\.\s", content.strip()):
                    om = re.match(r"(\d+)\.\s+(.*)", content.strip())
                    p.paragraph_format.left_indent = Inches(0.6)
                    add_runs(p, f"{om.group(1)}.  {om.group(2)}")
                else:
                    add_runs(p, content)
                    for r in p.runs:
                        r.italic = True
                        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                emit(p)
                i += 1
            continue

        # ordered list
        om = re.match(r"(\d+)\.\s+(.*)", stripped)
        if om:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_after = Pt(3)
            add_runs(p, f"{om.group(1)}.  ")
            add_runs(p, om.group(2))
            emit(p)
            i += 1
            continue

        # unordered list
        um = re.match(r"-\s+(.*)", stripped)
        if um:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_after = Pt(3)
            add_runs(p, "•  ")
            add_runs(p, um.group(1))
            emit(p)
            i += 1
            continue

        # plain paragraph
        p = doc.add_paragraph()
        add_runs(p, stripped)
        emit(p)
        i += 1

    # keep-together glue:
    # every block glues to the next, EXCEPT right after an answer box (box_end)
    # and the final block. This keeps each prompt stuck to its writing lines and
    # only allows a page break once a box has finished.
    for idx, (p, box_end) in enumerate(paras):
        last = idx == len(paras) - 1
        set_keep(p, with_next=not (box_end or last))

    doc.save(out_path)


def main():
    for md in sorted(glob.glob("*.md")):
        out = md[:-3] + ".docx"
        render(md, out)
        print("ok", out)


if __name__ == "__main__":
    main()
